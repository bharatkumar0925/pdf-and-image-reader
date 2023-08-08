from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import pdfencrypt
from ebooklib import epub
from docx import Document
from speech import tts


# Function to save the extracted text as PDF
def save_as_pdf(text, output_path, password=None):
    c = canvas.Canvas(output_path, letter)
    c.setFont("Helvetica", 12)
    lines = text.split('\n')
    y = 800
    for line in lines:
        c.drawString(50, y, line)
        y -= 20
    c.save()

    # add the functionality for password
    if password:
        encrypt = pdfencrypt.StandardEncryption(password, canPrint=0, canCopy=0, canModify=0)
        with open(output_path, 'rb') as f:
            data = f.read()
        with open(output_path, 'wb') as f:
            encrypt.encrypt(data, f)

# Function to save the extracted text as EPUB
def save_as_epub(text, output_path):
    book = epub.EpubBook()
    book.set_title("Extracted Text")
    book.add_author("Your Name")

    # Add a new chapter
    c = epub.EpubHtml(title="Chapter 1", file_name="chap_1.xhtml", lang="en")
    c.content = "<html><body><p>" + text.replace("\n", "<br>") + "</p></body></html>"
    book.add_item(c)

    # Define Table of Contents
    book.toc = (epub.Link("chap_1.xhtml", "Chapter 1", "chapter"),)
    book.spine = ['nav', c]

    # Save the EPUB
    epub.write_epub(output_path, book, {})

# Function to save the extracted text as DOCX
def save_as_docx(text, output_path, password='i am'):
    doc = Document()
    doc.add_heading('Extracted Text', level=1)
    doc.add_paragraph(text)
    doc.save(output_path)
    # Set password for the DOCX if provided
    if password:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = 0
        word.Documents.Open(output_path)
        doc = word.ActiveDocument
        doc.Password = password
        doc.SaveAs(output_path)
        doc.Close()
        word.Quit()

def save_as_txt(text, output_path):
    with open(output_path, 'w') as file:
        file.write(text)



def save_as_file(text, output_path, format='txt'):
    if format == 'mp3':
        tts(text, output_path, save_file=True)
    elif format == 'pdf':
        save_as_pdf(text, output_path)
    elif format == 'epub':
        save_as_epub(text, output_path)
    elif format=='docx':
        save_as_docx(text, output_path)
    else:
        save_as_txt(text, output_path)

